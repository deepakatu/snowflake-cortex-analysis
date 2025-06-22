from snowflake_utils import get_snowpark_session_from_config
import pprint as pprint
import time
import pandas as pd
import snowflake.snowpark.functions as F
from snowflake.cortex import complete
from snowflake.snowpark.functions import col, to_date, date_part, sum as Fsum, count as Fcount
from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading("Cortex Analysis Output", 0)

start_time = time.time()
pp = pprint.pprint

session = get_snowpark_session_from_config(connection_name="dev", config_path="config.toml")
reviews_df = pd.read_csv("tbvoc-reviews-clean.csv", delimiter="|")
print(reviews_df.head(10))

print("Access data completeness: by month")
snowpark_df = session.create_dataframe(reviews_df)
snowpark_df = snowpark_df.with_column("REVIEW_DATE", to_date(col("REVIEW_DATE")))
snowpark_df_temp = (
    snowpark_df.with_column("MONTH", date_part("month", col("REVIEW_DATE")))
    .group_by(["TRUCK_BRAND_NAME", "MONTH"])
    .agg(Fcount("*").alias("NUM_REVIEWS"))
)
reviews_pivot = (
    snowpark_df_temp.group_by("TRUCK_BRAND_NAME")
    .pivot("MONTH")
    .agg(Fsum("NUM_REVIEWS"))
    .sort("TRUCK_BRAND_NAME")
)
reviews_pivot.show(20)

print("Identify the top themes...")
reviews_str = "".join(map(str, reviews_df["REVIEW"].astype(str).tolist()))
prompt = f"""[INST]###Summarize the issues mentioned in following  \
food truck customer reviews with five concise bullet points  \
under 50 words each. Append to each bullet point the number of reviews that relate to that bullet point. \   
    Each bullet point also has a heading along with recommendations to remedy those issues.###### {reviews_str} [/INST]"""

query = f"""
    SELECT snowflake.cortex.complete('llama3.1-70b', $$ {prompt} $$) AS output
"""
result = session.sql(query).collect()[0]["OUTPUT"]
print(result)
doc.add_heading("Identify the top themes", level=1)
doc.add_paragraph(result)

print("Create a draft email that summarises all the themes")
email_prompt = f"""[INST]### Write a professional and friendly email that summarises the themes below and recommends prompt action. Less than 300 words.### {result} [/INST]"""
email_query = f"""
    SELECT snowflake.cortex.complete('llama3.1-70b', $$ {email_prompt} $$) AS output
"""
email = session.sql(email_query).collect()[0]["OUTPUT"]
print(email)
doc.add_heading("Draft Email Summary", level=1)
doc.add_paragraph(email)

print("Summarise reviews mentioning food quality")
food_quality_prompt = f"""[INST]###Summarize reviews mentioning food quality. Emphasise the negatives. Under 300 words.### {reviews_str} [/INST]"""
food_quality_query = f"""
    SELECT snowflake.cortex.complete('llama3.1-70b', $$ {food_quality_prompt} $$) AS output
"""
summary = session.sql(food_quality_query).collect()[0]["OUTPUT"]
print(summary)
doc.add_heading("Food Quality Review Summary", level=1)
doc.add_paragraph(summary)

print("Compare 'food quality' reviews for M1-3 vs M4-6")
compare_prompt = f"""[INST]###Compare reviews on food quality between months 1-3 and 4-6. Ignore sentiment. Under 200 words.### {reviews_str} [/INST]"""
compare_query = f"""
    SELECT snowflake.cortex.complete('llama3.1-70b', $$ {compare_prompt} $$) AS output
"""
comparison = session.sql(compare_query).collect()[0]["OUTPUT"]
print(comparison)
doc.add_heading("Food Quality Change (Month 1-3 vs 4-6)", level=1)
doc.add_paragraph(comparison)

total_time = round((time.time() - start_time) / 60, 2)
print(f"<<< END >>> {total_time} minutes\n")
doc.save("cortex_analysis_output.docx")
print("Word document saved: cortex_analysis_output.docx")
